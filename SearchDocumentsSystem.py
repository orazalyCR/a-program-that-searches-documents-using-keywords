import os

from docx import Document

from tkinter import *
from tkinter import filedialog

LIST_END = ""

match_items = []
match_word = ""

bg_color = "#3d6466"
bg_color_btn = "#377679"

# Открываем нужную директорию с документами
def openFileDialog():
    fname = filedialog.askdirectory()
    textbox_ofd.configure(state=NORMAL)
    textbox_ofd.delete(0,END)
    textbox_ofd.insert(0,fname)
    textbox_ofd.configure(state=DISABLED)
    folder = textbox_ofd.get()
    if not (folder == ""):
        file_names = os.listdir(folder)
        file_names = [file for file in file_names if file.endswith('.docx')]
        file_names = [os.path.join(folder, file) for file in file_names]
        list_docs.delete(0,END)
        for file in file_names:
            document = Document(file)
            file = os.path.basename(file)
            list_docs.insert(END,file)
        list_docs.insert(END,LIST_END)
    else:
        textbox_ofd.configure(state=NORMAL)
        textbox_ofd.insert(0, "Путь не указан. Выберите папку.")
        textbox_ofd.configure(state=DISABLED)

# Поиск по ключевым словам
def search():
    match_items = []
    list_docs.delete(0, END)
    
    count_hp = 0
    count_t = 0
    folder = textbox_ofd.get()
    file_names = os.listdir(folder)
    file_names = [file for file in file_names if file.endswith('.docx')]
    file_names = [os.path.join(folder, file) for file in file_names]
    
    match_word = textbox_search.get().lower()
    
    if not (match_word == ""):
        for file in file_names:
            document = Document(file)
            file = os.path.basename(file)
            
            for paragraph in document.paragraphs:
                paragraph.text = paragraph.text.lower()
                if match_word in paragraph.text:
                    count_hp = count_hp + 1

            for table in document.tables:
                for row in table.rows:
                    for i in range(len(row.cells)):
                        row.cells[i].text = row.cells[i].text.lower()
                        if match_word in row.cells[i].text:
                            count_t = count_t + 1
                
            match_items.append([file, "Заголовки и параграфы: " + str(count_hp), "Таблицы: " + str(count_t)])
            count_hp = 0
            count_t = 0
            
        for i in range(len(match_items)):
            if not ((match_items[i][1] == "Заголовки и параграфы: 0") and (match_items[i][2] == "Таблицы: 0")):
                list_docs.insert(END, [match_items[i][0], match_items[i][1], match_items[i][2]])
    
    else:
        folder = textbox_ofd.get()
        file_names = os.listdir(folder)
        file_names = [file for file in file_names if file.endswith('.docx')]
        file_names = [os.path.join(folder, file) for file in file_names]
        list_docs.delete(0,END)
        for file in file_names:
            document = Document(file)
            file = os.path.basename(file)
            list_docs.insert(END,file)

    list_docs.insert(END,LIST_END)
# Открытие выбранного документа    
def openDoc():
    f = textbox_ofd.get()
    tag = list_docs.get(list_docs.curselection())
    if not(tag == LIST_END):
        if not (isinstance(tag,tuple)):
            text = f + '/' + tag
            os.startfile(text)
            f       = ""
            tag     = ""
            text    = ""
        else:
            text = f + '/' + tag[0]
            os.startfile(text)
            f       = ""
            tag     = ""
            text    = ""
    
# Создание окна
root = Tk()
root.title("Поиск документов")
root.geometry("640x480")

# Фреймы нужны для создания блоков элементов
frm1 = Frame(root, bg=bg_color)
frm1.place(relx=0,rely=0,relwidth=1.0,relheight=0.1, anchor="nw")

frm2 = Frame(root,  bg="#00F3DD")
frm2.place(relx=0,rely=0.1,relwidth=1.0,relheight=0.9, anchor="nw")

frm3 = Frame(root, bg="#8E2693")
frm3.place(relx=0,rely=0.9,relwidth=1.0,relheight=0.1, anchor="nw")

# Текстовое поле для выбора папки
textbox_ofd = Entry(
    frm1, 
    bg = bg_color,
    fg = "white",
    font= ('Arial', 16)
)
textbox_ofd.configure(state=NORMAL)
textbox_ofd.insert(0, "Путь не указан.")
textbox_ofd.configure(state=DISABLED)

textbox_ofd.place(x=0,y=0,anchor="nw",relwidth=0.8,relheight=1.0)

# Кнопка для выбора папки
btn_ofd = Button(
    frm1,
    text = "Открыть",
    font= ('Arial', 16),
    bg = bg_color_btn,
    fg = "white",
    cursor="hand2",
    activebackground="#badee2",
    activeforeground="black",
    command=lambda:openFileDialog()
)

btn_ofd.place(relx=1.0,rely=0,anchor="ne",relwidth=0.2,relheight=1.0)

# Список документов
list_docs = Listbox(
    frm2,
    font= ('Arial', 16),
    bg = "#28393a",
    fg = "white",
    selectmode=SINGLE
    )

list_docs.place(x=0,y=0,anchor="nw",relwidth=0.98,relheight=0.87)
list_docs.bind('<Double-Button>', lambda x:openDoc())
list_docs.insert(0,LIST_END)

# Элемент для пролистывания списка по координат Y
scrollbar_y = Scrollbar(frm2, orient="vertical")
scrollbar_y.config(command=list_docs.yview)

list_docs.config(yscrollcommand=scrollbar_y.set)

scrollbar_y.place(relx=1.0,rely=0,relwidth=0.02,relheight=0.89, anchor="ne")

# Элемент для пролистывания списка по координат X
scrollbar_x = Scrollbar(frm2, orient="horizontal")
scrollbar_x.config(command=list_docs.xview)

list_docs.config(xscrollcommand=scrollbar_x.set)

scrollbar_x.place(relx=0,rely=0.89,anchor="sw",relwidth=0.98,relheight=0.03)

# Текстовое поле для поиска по ключевым словам
textbox_search = Entry(
    frm3, 
    bg = bg_color,
    fg = "white",
    font= ('Arial', 16)
)

textbox_search.place(x=0,y=0,anchor="nw",relwidth=0.8,relheight=1.0)

# Кнопка для поиска по ключевым словам
btn_search = Button(
    frm3,
    text = "Найти",
    font= ('Arial', 16),
    bg = bg_color_btn,
    fg = "white",
    cursor="hand2",
    activebackground="#badee2",
    activeforeground="black",
    command=lambda:search()
)

btn_search.place(relx=1.0,rely=0,anchor="ne",relwidth=0.2,relheight=1.0)

# Запуск программы
root.mainloop()